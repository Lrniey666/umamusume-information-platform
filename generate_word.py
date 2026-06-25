"""
將論文_程式執行過程與結果.md 轉為符合台灣論文格式的 .docx
"""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── 顏色常數 ──────────────────────────────────────────
C_DARK      = RGBColor(0x1A, 0x1A, 0x2E)   # 封面深藍
C_PURPLE    = RGBColor(0x6A, 0x35, 0xB4)   # 紫色強調
C_GOLD      = RGBColor(0xC9, 0xA0, 0x27)   # 金色副標
C_TABLE_H   = RGBColor(0x4A, 0x23, 0x8A)   # 表頭背景
C_TABLE_H2  = RGBColor(0x6B, 0x3F, 0xB5)   # 表頭第二色
C_GRAY      = RGBColor(0xF5, 0xF5, 0xF7)   # 表格奇數列底色
C_CODE_BG   = RGBColor(0xF0, 0xF0, 0xF5)   # 程式碼底色
C_BLACK     = RGBColor(0x1A, 0x1A, 0x1A)
C_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
C_FIG       = RGBColor(0x55, 0x55, 0x88)   # 圖片說明色

# ── helper: 設定儲存格底色 ────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}'
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{side}')
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), '4')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), 'CCCCCC')
        tcBorders.append(tag)
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcBorders)

def add_run_bold_inline(para, text: str, base_font='標楷體', base_size=12):
    """處理段落內 **粗體** 與 `行內程式碼`，其餘為普通文字。"""
    pattern = re.compile(r'\*\*(.+?)\*\*|`([^`]+)`')
    last = 0
    for m in pattern.finditer(text):
        # 前段普通文字
        if m.start() > last:
            r = para.add_run(text[last:m.start()])
            r.font.name = base_font
            r.font.size = Pt(base_size)
            r.font.color.rgb = C_BLACK
        if m.group(1):  # **bold**
            r = para.add_run(m.group(1))
            r.bold = True
            r.font.name = base_font
            r.font.size = Pt(base_size)
            r.font.color.rgb = C_BLACK
        elif m.group(2):  # `code`
            r = para.add_run(m.group(2))
            r.font.name = 'Consolas'
            r.font.size = Pt(10)
            r.font.color.rgb = C_PURPLE
        last = m.end()
    # 剩餘文字
    if last < len(text):
        r = para.add_run(text[last:])
        r.font.name = base_font
        r.font.size = Pt(base_size)
        r.font.color.rgb = C_BLACK

def set_para_spacing(para, before=0, after=0, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line:
        pf.line_spacing = Pt(line)

def set_page_margins(doc, top=2.5, bottom=2.5, left=3.0, right=2.5):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)

# ── 建立 Document ─────────────────────────────────────
doc = Document()
set_page_margins(doc)

# 預設字體（Normal style）
style = doc.styles['Normal']
style.font.name = '標楷體'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

# ══════════════════════════════════════════════════════
# 封面頁
# ══════════════════════════════════════════════════════
def add_cover(doc):
    # 上方留白
    for _ in range(4):
        p = doc.add_paragraph()
        set_para_spacing(p, 0, 0)

    # 主標題
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, 0, 6)
    r = p.add_run('賽馬娘公眾輿情分析平台')
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = C_DARK
    r.font.name = '標楷體'

    # 副標題
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, 0, 30)
    r = p.add_run('程式執行過程及結果')
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = C_PURPLE
    r.font.name = '標楷體'

    # 分隔線
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, 0, 20)
    r = p.add_run('─' * 30)
    r.font.color.rgb = C_GOLD
    r.font.size = Pt(14)

    # 資訊區
    for label, value in [
        ('平台版本', 'v0.6.51-alpha'),
        ('技術棧', 'Django 5.2 LTS · Gemini 3.5 Flash · Claude Sonnet 4.6'),
        ('撰寫日期', '2026 年 6 月 25 日'),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p, 0, 4)
        r = p.add_run(f'{label}：')
        r.bold = True
        r.font.name = '標楷體'
        r.font.size = Pt(12)
        r = p.add_run(value)
        r.font.name = '標楷體'
        r.font.size = Pt(12)

    # 分頁
    doc.add_page_break()

add_cover(doc)

# ══════════════════════════════════════════════════════
# 摘要頁
# ══════════════════════════════════════════════════════
def add_abstract(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, 0, 12)
    r = p.add_run('摘　要')
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = '標楷體'
    r.font.color.rgb = C_DARK

    abstract_text = (
        '本報告旨在說明「賽馬娘公眾輿情分析平台」（Uma Musume Information Platform）之程式執行過程及各功能模組的實際執行結果。'
        '本平台以 Django 5.2 LTS 為核心框架，整合多來源爬蟲資料管線、大型語言模型（LLM）分析、Agentic AI 代理人架構、RAG 向量知識庫、'
        'APScheduler 自動排程，以及 Docker Compose 容器化部署，針對賽馬娘 Pretty Derby 遊戲社群之輿情數據進行全方位分析與呈現。\n\n'
        '本平台共實作 28 個功能模組，資料涵蓋巴哈姆特哈啦板、Bilibili BWIKI、ETtoday、UDN 及 Gamme 等五大來源，'
        '並額外串接 YouTube Data API v3 及 Discord Bot 自動推播系統，提供即時輿情監控、情感分析、Agentic AI 多工具問答、'
        '以及 3D 虛擬客服互動等功能，為課程期末專案中技術完整性與主題一致性兼備之完整輿情分析解決方案。'
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.85)
    set_para_spacing(p, 0, 8, line=22)
    r = p.add_run(abstract_text)
    r.font.name = '標楷體'
    r.font.size = Pt(12)
    r.font.color.rgb = C_BLACK

    p = doc.add_paragraph()
    set_para_spacing(p, 8, 0)
    r = p.add_run('關鍵詞：')
    r.bold = True
    r.font.name = '標楷體'
    r.font.size = Pt(12)
    r = p.add_run('公眾輿論分析、Django、大型語言模型、Agentic AI、RAG、賽馬娘')
    r.font.name = '標楷體'
    r.font.size = Pt(12)

    doc.add_page_break()

add_abstract(doc)

# ══════════════════════════════════════════════════════
# 目錄頁
# ══════════════════════════════════════════════════════
def add_toc(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, 0, 16)
    r = p.add_run('目　錄')
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = '標楷體'
    r.font.color.rgb = C_DARK

    toc_items = [
        ('摘要', ''),
        ('第一章　系統架構概述', ''),
        ('　1.1　專案背景與目標', ''),
        ('　1.2　技術選型', ''),
        ('　1.3　系統架構圖', ''),
        ('　1.4　功能模組一覽', ''),
        ('第二章　執行環境與部署流程', ''),
        ('　2.1　執行環境需求', ''),
        ('　2.2　環境建置步驟', ''),
        ('　2.3　Docker Compose 容器化啟動', ''),
        ('第三章　資料管線執行過程', ''),
        ('　3.1　資料來源概述', ''),
        ('　3.2　資料管線流程', ''),
        ('第四章　核心功能模組執行結果', ''),
        ('　4.1　首頁：角色人氣 PK', ''),
        ('　4.2　關鍵詞聲量分析', ''),
        ('　4.3　情感分析', ''),
        ('　4.4　AI 雙模型分析報告', ''),
        ('　4.5　巴哈留言情感排程儀表板', ''),
        ('　4.6　熱門關鍵詞排行', ''),
        ('　4.7　公告儀表板', ''),
        ('　4.8　YouTube 影片情感儀表板', ''),
        ('第五章　Agentic AI 與 RAG 知識庫執行結果', ''),
        ('　5.1　Agentic AI 馬娘分析助理', ''),
        ('　5.2　LangChain ReAct Agent', ''),
        ('　5.3　LangGraph 狀態圖 Agent', ''),
        ('　5.4　RAG 馬娘知識庫問答', ''),
        ('　5.5　Agentic RAG 雙工具問答', ''),
        ('　5.6　全站 3D 虛擬客服「成田路」', ''),
        ('第六章　排程系統與 Discord Bot 執行結果', ''),
        ('　6.1　APScheduler 自動排程系統', ''),
        ('　6.2　Discord Bot 自動新聞推播系統', ''),
        ('第七章　情報站控制台執行結果', ''),
        ('　7.1　控制台總覽', ''),
        ('　7.2　子功能頁面', ''),
        ('第八章　系統部署與容器化執行結果', ''),
        ('　8.1　Docker Compose 架構', ''),
        ('　8.2　部署執行結果', ''),
        ('　8.3　靜態檔案與媒體服務確認', ''),
        ('第九章　執行成效彙整', ''),
        ('　9.1　功能完成度彙整', ''),
        ('　9.2　系統性能指標', ''),
        ('　9.3　資料規模', ''),
        ('　9.4　平台架構亮點', ''),
        ('附錄 A　目錄結構一覽', ''),
        ('附錄 B　主要 API 端點一覽', ''),
        ('附錄 C　環境變數說明', ''),
    ]

    for title, page in toc_items:
        p = doc.add_paragraph()
        set_para_spacing(p, 0, 3, line=18)
        is_chapter = title.startswith('第') or title.startswith('摘') or title.startswith('附')
        r = p.add_run(title)
        r.font.name = '標楷體'
        r.font.size = Pt(12 if is_chapter else 11)
        r.bold = is_chapter
        r.font.color.rgb = C_DARK if is_chapter else C_BLACK

    doc.add_page_break()

add_toc(doc)

# ══════════════════════════════════════════════════════
# 內文工具函數
# ══════════════════════════════════════════════════════
def add_chapter_heading(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, 18, 10)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = '標楷體'
    r.font.color.rgb = C_DARK

def add_section_heading(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, 12, 6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = '標楷體'
    r.font.color.rgb = C_PURPLE

def add_subsection_heading(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, 8, 4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = '標楷體'
    r.font.color.rgb = C_DARK

def add_body(doc, text, indent_first=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(0.85)
    set_para_spacing(p, 0, 6, line=22)
    add_run_bold_inline(p, text)

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph()
    set_para_spacing(p, 0, 3, line=20)
    indent = Cm(0.8 + level * 0.6)
    p.paragraph_format.left_indent  = indent
    p.paragraph_format.first_line_indent = Cm(-0.4)
    r = p.add_run('• ' if level == 0 else '◦ ')
    r.font.name = 'Arial'
    r.font.size = Pt(11)
    r.font.color.rgb = C_PURPLE
    add_run_bold_inline(p, text, base_size=12)

def add_numbered(doc, num, text):
    p = doc.add_paragraph()
    set_para_spacing(p, 0, 3, line=20)
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.8)
    r = p.add_run(f'{num}. ')
    r.bold = True
    r.font.name = '標楷體'
    r.font.size = Pt(12)
    r.font.color.rgb = C_PURPLE
    add_run_bold_inline(p, text, base_size=12)

def add_figure_placeholder(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, 10, 10)
    # 框線段落（用 border 模擬）
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), '9B72D0')
        pBdr.append(el)
    pPr.append(pBdr)
    r = p.add_run(f'  {text}  ')
    r.font.name = '標楷體'
    r.font.size = Pt(11)
    r.font.color.rgb = C_FIG
    r.italic = True

def add_code_block(doc, code_text):
    """灰底等寬字程式碼區塊（每行一個段落）"""
    lines = code_text.strip().split('\n')
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        if i == 0:
            set_para_spacing(p, 6, 0)
        elif i == len(lines) - 1:
            set_para_spacing(p, 0, 6)
        else:
            set_para_spacing(p, 0, 0)
        # 灰底
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F0F0F5')  # code block bg
        pPr.append(shd)
        r = p.add_run(line if line else ' ')
        r.font.name = 'Consolas'
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x2D, 0x2D, 0x6E)

def add_table_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, 8, 4)
    r = p.add_run(text)
    r.bold = True
    r.font.name = '標楷體'
    r.font.size = Pt(12)
    r.font.color.rgb = C_DARK

def styled_table(doc, headers, rows, col_widths=None):
    """建立帶格式的表格"""
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 表頭
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        set_cell_bg(cell, C_TABLE_H)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(h)
        r.bold = True
        r.font.name = '標楷體'
        r.font.size = Pt(11)
        r.font.color.rgb = C_WHITE

    # 資料列
    for ri, row_data in enumerate(rows):
        row_cells = table.rows[ri+1].cells
        bg = C_GRAY if ri % 2 == 0 else C_WHITE
        for ci, val in enumerate(row_data):
            cell = row_cells[ci]
            set_cell_bg(cell, bg)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            # 程式碼樣式（cell 中的 code 片段）
            if '`' in val or val.startswith('/') or val.startswith('app_'):
                r = p.add_run(val)
                r.font.name = 'Consolas'
                r.font.size = Pt(10)
                r.font.color.rgb = C_PURPLE
            else:
                add_run_bold_inline(p, val, base_size=11)

    # 欄寬
    if col_widths:
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                if ci < len(col_widths):
                    cell.width = Cm(col_widths[ci])

    # 段落間距
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table

# ══════════════════════════════════════════════════════
# 第一章　系統架構概述
# ══════════════════════════════════════════════════════
add_chapter_heading(doc, '第一章　系統架構概述')

add_section_heading(doc, '1.1　專案背景與目標')
add_body(doc, '本專案以課程期末作業「公眾輿論分析（Public Opinion Analysis, POA）」為背景，選用**賽馬娘 Pretty Derby** 遊戲社群作為分析主題，以先前作業成果（w12 umamusume-llm-report）為基底，擴展為涵蓋資料爬取、情感分析、LLM 報告、Agentic AI、RAG 知識庫，以及 Docker 部署之完整平台。')
add_body(doc, '本平台須同時滿足兩個層次之要求：', indent_first=False)
add_numbered(doc, 1, '**技術完整性**：涵蓋課程 w12–w16 所有核心技術（LLM 報告、留言情感分析、Agentic AI、RAG、Docker 部署）')
add_numbered(doc, 2, '**主題一致性**：全站以馬娘遊戲輿情為核心，資料集、分析邏輯及 AI 提示詞（Prompt）均圍繞遊戲社群討論')

add_section_heading(doc, '1.2　技術選型')
add_body(doc, '表 1.1 列出本平台所採用之核心技術元件及其版本。')
add_table_caption(doc, '表 1.1　核心技術選型一覽')
styled_table(doc,
    ['層次', '技術元件', '採用版本', '說明'],
    [
        ['Web 框架', 'Django + Gunicorn', '5.2.15 LTS + ≥23.0', 'LTS 支援至 2028 年 4 月'],
        ['反向代理', 'Nginx', '1.26 stable', 'Alpine 映像，輕量高效'],
        ['資料庫', 'SQLite（開發）/ PostgreSQL（正式）', '— / postgres:16-alpine', 'ORM 統一操作'],
        ['LLM（Gemini）', 'google-genai', '2.9.0', '取代已棄用 google-generativeai'],
        ['Gemini 模型', 'gemini-3.5-flash', '—', '2.0-flash 已於 2026-06-01 停服'],
        ['LLM（Claude）', 'anthropic', '0.111.0', '雙模型對照分析'],
        ['Claude 模型', 'claude-sonnet-4-6', '—', '量產首選，平衡效能與成本'],
        ['向量搜尋', 'faiss-cpu', '1.14.2', 'RAG 語意檢索核心'],
        ['排程', 'APScheduler + django-apscheduler', '3.11.2 + 0.7.0', '自動爬取與情感標記'],
        ['前端圖表', 'Chart.js', '4.x（CDN）', '聲量趨勢與情感分布圖表'],
        ['CSS 框架', 'Bootstrap', '5.3（CDN）', '響應式設計'],
        ['Agent 框架', 'LangChain + LangGraph', '≥1.3.0 + ≥0.2.0', 'Agentic AI 選用項'],
        ['Discord Bot', 'discord.py', '2.4.0', '自動新聞推播系統'],
        ['容器化', 'Docker + Docker Compose v2', 'Engine 27.x', '生產環境部署'],
    ],
    col_widths=[3.0, 4.5, 4.0, 4.5]
)

add_section_heading(doc, '1.3　系統架構圖')
add_body(doc, '本平台整體架構如圖 1.1 所示，採用 Nginx + Gunicorn + Django 三層架構，外部透過 Docker Compose 進行容器化管理。')
add_figure_placeholder(doc, '【圖 1.1　系統架構總覽圖】')
add_body(doc, '系統各層說明如下：', indent_first=False)
for item in [
    '**Nginx（反向代理層）**：監聽 Port 80，負責靜態檔案（`/static/`）的直接服務，以及所有動態請求的反向代理至 Gunicorn。',
    '**Django + Gunicorn（應用層）**：包含 28 個功能 App，處理路由分派、業務邏輯、ORM 資料存取，以及與外部 API（Gemini、Claude、YouTube、Discord）的整合。',
    '**SQLite / PostgreSQL（資料層）**：儲存所有結構化資料（NewsData、Article/Comment/ArticleEmotion、YouTubeVideo/Comment、DiscordMessage 等）。',
    '**FAISS 向量索引（知識庫層）**：`app_rag_uma/index/` 預建持久化向量索引，供 RAG 問答模組語意檢索使用。',
    '**離線資料管線（Pipeline）**：於部署前執行，將五個來源的原始爬蟲資料逐步轉換為可分析的結構化資料並匯入資料庫。',
]:
    add_bullet(doc, item)

add_section_heading(doc, '1.4　功能模組一覽')
add_body(doc, '本平台共實作 28 個功能模組，依功能類型分組如表 1.2 所示。')
add_table_caption(doc, '表 1.2　功能模組分組一覽')
styled_table(doc,
    ['功能類別', 'App 模組', '主要功能'],
    [
        ['基礎輿情分析', 'app_user_keyword', '自訂關鍵詞聲量分析（CSV 版）'],
        ['', 'app_user_keyword_association', '全文關聯分析'],
        ['', 'app_user_keyword_sentiment', '關鍵詞情感分析'],
        ['', 'app_user_keyword_db', '全文資料庫搜尋（ORM）'],
        ['', 'app_correlation_analysis', '關鍵詞相關性分析'],
        ['排行統計', 'app_uma_top_keyword', '熱門關鍵詞排行（5 類別）'],
        ['', 'app_uma_top_character', '熱門角色排行（5 類別）'],
        ['特色功能', 'app_character_pk', '首頁角色人氣 PK 比較'],
        ['', 'app_comment_sentiment', '留言情感排程儀表板'],
        ['LLM 分析', 'app_user_keyword_llm_report', '雙模型 AI 報告（Gemini + Claude）'],
        ['Agentic AI', 'app_agent_uma', 'Agentic AI（6 種 Function Calling）'],
        ['', 'app_agent_langchain', 'LangChain ReAct Agent'],
        ['', 'app_agent_langgraph', 'LangGraph StateGraph Agent'],
        ['', 'app_rag_agent', 'Agentic RAG（雙工具 Agent）'],
        ['RAG 知識庫', 'app_rag_uma', 'RAG + FAISS 向量問答'],
        ['資料模型', 'app_uma_news', '遊戲公告資料模型'],
        ['', 'app_uma_comments', '巴哈留言資料模型'],
        ['', 'app_dashboard', '公告列表儀表板'],
        ['外部服務', 'app_youtube_uma', 'YouTube 影片情感儀表板'],
        ['', 'app_discord_bot', 'Discord Bot 自動新聞推播'],
        ['管理後台', 'app_crawler_admin', '情報站控制台'],
        ['介紹頁', 'app_poa_introduction', '平台介紹頁'],
        ['', 'app_course_intro', '課程技術說明頁'],
    ],
    col_widths=[3.5, 5.5, 7.0]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 第二章　執行環境與部署流程
# ══════════════════════════════════════════════════════
add_chapter_heading(doc, '第二章　執行環境與部署流程')

add_section_heading(doc, '2.1　執行環境需求')
add_body(doc, '本平台可於以下兩種環境執行：', indent_first=False)
add_subsection_heading(doc, '（一）本機開發環境')
for item in [
    'Python 3.12',
    '相依套件：依 `requirements.txt` 安裝',
    '資料庫：SQLite（`db.sqlite3`）',
    '環境變數：`.env` 檔案（含 Gemini API Key、Anthropic API Key、YouTube API Key、Discord Bot Token）',
]:
    add_bullet(doc, item)

add_subsection_heading(doc, '（二）容器化正式環境')
for item in [
    'Docker Engine 27.x + Docker Compose v2',
    '服務容器：`web-poa`（Django + Gunicorn）、`discord-bot`（Discord Bot）、`nginx`（反向代理）',
    '執行埠：對外 Port 80',
]:
    add_bullet(doc, item)

add_section_heading(doc, '2.2　環境建置步驟')
add_subsection_heading(doc, '步驟一：複製專案並安裝相依套件')
add_code_block(doc, 'git clone <repository-url>\ncd umamusume-information-platform\npip install -r requirements.txt')

add_subsection_heading(doc, '步驟二：設定環境變數')
add_body(doc, '複製 `.env.example` 為 `.env`，填入各 API 金鑰：')
add_code_block(doc, 'GEMINI_API_KEY=<your-gemini-api-key>\nANTHROPIC_API_KEY=<your-anthropic-api-key>\nYOUTUBE_API_KEY=<your-youtube-api-key>\nDISCORD_BOT_TOKEN=<your-discord-bot-token>\nUMA_CHAT_MODEL=gemini-3.5-flash')

add_subsection_heading(doc, '步驟三：執行資料庫遷移')
add_code_block(doc, 'python manage.py migrate')

add_subsection_heading(doc, '步驟四：啟動本機開發伺服器')
add_code_block(doc, 'python manage.py runserver')
add_body(doc, '執行成功後，開啟瀏覽器前往 `http://localhost:8000` 即可存取平台。')
add_figure_placeholder(doc, '【圖 2.1　本機開發伺服器啟動成功畫面（終端機輸出）】')

add_section_heading(doc, '2.3　Docker Compose 容器化啟動')
add_body(doc, '在正式環境或課堂展示情境下，以 Docker Compose 一鍵啟動：')
add_code_block(doc, 'docker compose build\ndocker compose up -d')
add_body(doc, '`entrypoint.sh` 會在容器首次啟動時自動執行以下初始化作業：', indent_first=False)
for i, item in enumerate([
    '執行 `python manage.py migrate` — 建立資料庫結構',
    '執行 `python manage.py collectstatic --noinput` — 收集靜態檔案',
    '匯入預處理資料（idempotent 保護，重複啟動不重複匯入）',
    '建立 RAG FAISS 向量索引（如不存在）',
    '啟動 Gunicorn WSGI 伺服器',
], 1):
    add_numbered(doc, i, item)
add_figure_placeholder(doc, '【圖 2.2　Docker Compose 啟動成功畫面（`docker compose ps` 輸出）】')

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 第三章　資料管線執行過程
# ══════════════════════════════════════════════════════
add_chapter_heading(doc, '第三章　資料管線執行過程')

add_section_heading(doc, '3.1　資料來源概述')
add_body(doc, '本平台整合五個資料來源，如表 3.1 所示。')
add_table_caption(doc, '表 3.1　資料來源一覽')
styled_table(doc,
    ['來源代碼', '平台', '資料類型', '爬蟲方式'],
    [
        ['bahamut', '巴哈姆特哈啦板（bsn=34421）', '討論串與留言', '巴哈 API 爬蟲'],
        ['bilibili', 'Bilibili BWIKI 官方公告', '遊戲公告文章', 'HTML 解析'],
        ['ettoday', 'ETtoday 遊戲新聞', '新聞報導', 'HTML 解析'],
        ['udn', 'UDN 遊戲版', '新聞報導', 'HTML 解析'],
        ['gamme', 'Gamme 遊戲電影', '遊戲資訊', 'HTML 解析'],
    ],
    col_widths=[3.0, 5.0, 4.0, 4.0]
)

add_section_heading(doc, '3.2　資料管線流程')
add_body(doc, '本平台採用離線資料管線（Pipeline）架構，如圖 3.1 所示，在部署前完成所有資料處理。')
add_figure_placeholder(doc, '【圖 3.1　資料管線流程圖】')

add_subsection_heading(doc, '步驟一：原始資料爬取')
add_body(doc, '各爬蟲腳本執行後，將資料統一輸出至 `data/raw/` 目錄，採用標準化命名格式：')
add_code_block(doc, 'data/raw/bahamut_uma_raw.csv\ndata/raw/bilibili_uma_raw.csv\ndata/raw/ettoday_uma_raw.csv\ndata/raw/udn_uma_raw.csv\ndata/raw/gamme_uma_raw.csv')
add_body(doc, '原始 CSV 核心欄位包括：`title`、`content`、`url`、`published_at`、`source`、`category`。')

add_subsection_heading(doc, '步驟二：資料前處理（preprocess.py）')
add_body(doc, '執行 `python pipeline/preprocess.py`，將五個來源的原始 CSV 合併、清洗、斷詞，輸出至 `data/processed/uma_combined_tokenized.csv`。')
add_body(doc, '此步驟主要作業包含：移除重複資料、正規化日期格式（YYYY-MM-DD）、結巴（jieba）中文斷詞、標記公告類別。')
add_figure_placeholder(doc, '【圖 3.2　preprocess.py 執行終端機輸出】')

add_subsection_heading(doc, '步驟三：情感標記（label_sentiment.py）')
add_body(doc, '執行 `python pipeline/label_sentiment.py`，呼叫 Gemini API（gemini-3.5-flash）對每筆資料進行情感分類（正面 / 負面 / 中性），輸出至 `data/processed/uma_news_preprocessed.csv`。核心欄位新增：`sentiment`（正面/負面/中性）、`sentiment_score`（0.0–1.0）。')
add_figure_placeholder(doc, '【圖 3.3　label_sentiment.py 執行過程（逐批呼叫 Gemini API）】')

add_subsection_heading(doc, '步驟四：資料匯入資料庫')
add_body(doc, '執行 `python scripts/import_uma_data.py`，將 `uma_news_preprocessed.csv` 匯入 Django ORM 的 `NewsData` 資料表。')

add_subsection_heading(doc, '步驟五：統計資料生成')
add_body(doc, '執行以下腳本，生成熱門關鍵詞與熱門角色的預處理統計 CSV：')
add_code_block(doc, 'python scripts/generate_topkey_csv.py\npython scripts/generate_top_character_csv.py')
add_figure_placeholder(doc, '【圖 3.4　情報站控制台資料統計卡片，顯示各來源匯入筆數】')

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 第四章　核心功能模組執行結果
# ══════════════════════════════════════════════════════
add_chapter_heading(doc, '第四章　核心功能模組執行結果')

# 4.1
add_section_heading(doc, '4.1　首頁：角色人氣 PK（app_character_pk）')
add_subsection_heading(doc, '功能說明')
add_body(doc, '使用者於首頁（`http://localhost:8000/`）看到馬娘角色卡片牆，選取多個角色後點擊「比較」，系統以雷達圖與長條圖呈現各角色在巴哈版的討論聲量與情感多維對比。')
add_subsection_heading(doc, '執行過程')
for i, item in enumerate([
    '使用者於首頁卡片牆勾選欲比較的角色（可複選）',
    '點擊「開始比較」按鈕，前端以 AJAX 呼叫後端 API',
    '後端查詢 `NewsData` 與 `Article` 資料表，計算各角色出現次數及情感分布',
    '回傳 JSON 資料，前端以 Chart.js 渲染雷達圖（多維對比）與長條圖（聲量排行）',
], 1):
    add_numbered(doc, i, item)
add_figure_placeholder(doc, '【圖 4.1　首頁角色卡片牆畫面】')
add_figure_placeholder(doc, '【圖 4.2　角色人氣 PK 雷達圖與長條圖比較結果】')

# 4.2
add_section_heading(doc, '4.2　關鍵詞聲量分析（app_user_keyword）')
add_subsection_heading(doc, '功能說明')
add_body(doc, '使用者於聲量頁面（`/user-keyword/`）輸入關鍵字（如「卡池」、「活動」、「限定」），選擇類別與時間範圍，查詢結果以長條圖與折線圖展示出現次數與時間趨勢。')
add_subsection_heading(doc, '執行過程')
for i, item in enumerate([
    '使用者輸入關鍵字並選擇查詢參數',
    '後端讀取 `data/processed/uma_combined_tokenized.csv`，依斷詞結果過濾符合條件之資料列',
    '以 Pandas 彙整週維度聲量數據',
    '前端以 Chart.js 渲染長條圖（各類別分布）與折線圖（時間趨勢）',
], 1):
    add_numbered(doc, i, item)
add_figure_placeholder(doc, '【圖 4.3　關鍵詞「卡池」聲量分析長條圖與趨勢折線圖】')

# 4.3
add_section_heading(doc, '4.3　情感分析（app_user_keyword_sentiment）')
add_subsection_heading(doc, '功能說明')
add_body(doc, '使用者於情感頁面（`/user-keyword-sentiment/`）輸入關鍵字，查看相關公告的情感分布（正面 / 負面 / 中性比例）及每日趨勢折線圖。')
add_subsection_heading(doc, '執行過程')
for i, item in enumerate([
    '使用者輸入關鍵字',
    '後端從 `uma_news_preprocessed.csv` 篩選含關鍵字之資料列',
    '計算正面 / 負面 / 中性各佔比例及每日情感均值',
    '前端以圓餅圖呈現整體情感分布，折線圖呈現逐日情感趨勢',
], 1):
    add_numbered(doc, i, item)
add_figure_placeholder(doc, '【圖 4.4　情感分析圓餅圖（正/負/中性比例）與逐日趨勢折線圖】')

# 4.4
add_section_heading(doc, '4.4　AI 雙模型分析報告（app_user_keyword_llm_report）')
add_subsection_heading(doc, '功能說明')
add_body(doc, '使用者輸入關鍵字並選擇 AI 模型（Gemini 3.5 Flash 或 Claude Sonnet 4.6），系統自動整合聲量與情感數據組成提示詞（Prompt），呼叫對應 LLM 並在 30–40 秒內回傳至少 500 字、以 Markdown 排版的繁體中文分析報告。')
add_subsection_heading(doc, '執行過程')
add_numbered(doc, 1, '使用者輸入關鍵字並於下拉選單選擇模型')
add_numbered(doc, 2, '後端查詢聲量數據與情感統計，組合 Prompt：')
add_code_block(doc, '以下是關於「{keyword}」的分析數據：\n- 聲量：{volume_stats}\n- 情感分布：正面 {pos}%、負面 {neg}%、中性 {neu}%\n請以繁體中文撰寫一份至少 500 字的輿情分析報告，包含摘要、趨勢分析、建議與總結。')
add_numbered(doc, 3, '依選擇的模型呼叫對應 API：')
add_bullet(doc, '**Gemini 3.5 Flash**：透過 `google-genai 2.9.0` SDK，呼叫 `client.models.generate_content()`', level=1)
add_bullet(doc, '**Claude Sonnet 4.6**：透過 `anthropic 0.111.0` SDK，呼叫 `client.messages.create()`', level=1)
add_numbered(doc, 4, '後端接收回傳的 Markdown 文字，經 `Marked.js` 於前端渲染為格式化報告')
add_figure_placeholder(doc, '【圖 4.5　選擇 Gemini 3.5 Flash 生成的分析報告（Markdown 渲染畫面）】')
add_figure_placeholder(doc, '【圖 4.6　選擇 Claude Sonnet 4.6 生成的分析報告（雙模型對照）】')
add_subsection_heading(doc, '執行成效')
for item in [
    'Gemini 3.5 Flash 平均回應時間：約 15–25 秒',
    'Claude Sonnet 4.6 平均回應時間：約 20–35 秒',
    '生成報告長度：穩定達 500 字以上，含摘要、關鍵詞、趨勢分析、建議、總結五個段落',
]:
    add_bullet(doc, item)

# 4.5
add_section_heading(doc, '4.5　巴哈留言情感排程儀表板（app_comment_sentiment）')
add_subsection_heading(doc, '功能說明')
add_body(doc, '使用者開啟留言情感儀表板（`/comment-sentiment/`），查看最新爬取之巴哈馬娘版貼文列表，以及每篇貼文的 AI 情感分析結果，包含文章情感分數與留言情緒六維度（歡呼 / 開心 / 混亂 / 傻眼 / 憤怒 / 難過）的圓餅圖。')
add_subsection_heading(doc, '（一）資料爬取（scrape_bahamut）')
add_code_block(doc, 'python manage.py scrape_bahamut')
add_body(doc, '執行後從巴哈姆特哈啦板（bsn=34421）爬取最新貼文與留言，儲存至 `Article` 與 `Comment` 資料表。')
add_subsection_heading(doc, '（二）情感分析排程（APScheduler）')
for item in [
    '爬蟲任務：每 1 小時執行一次',
    '情感分析：每日 02:00 執行 `python manage.py analyze_sentiment`，呼叫 Gemini API 對每篇貼文及其留言進行六維度情緒標記',
]:
    add_bullet(doc, item)
add_figure_placeholder(doc, '【圖 4.7　巴哈留言情感排程儀表板（貼文列表與情緒圓餅圖）】')

# 4.6
add_section_heading(doc, '4.6　熱門關鍵詞排行（app_uma_top_keyword）')
add_body(doc, '使用者於熱門關鍵詞頁面（`/uma-top-keyword/`）選擇公告類別與 Top-K 數量，系統從預處理詞頻統計 CSV 直接讀取，以橫式長條圖呈現高頻詞彙排行，頁面回應速度快（< 1 秒）。')
add_figure_placeholder(doc, '【圖 4.8　各類別熱門關鍵詞橫式長條圖（以「卡池」類別 Top 30 為例）】')

# 4.7
add_section_heading(doc, '4.7　公告儀表板（app_dashboard）')
add_body(doc, '使用者於公告儀表板（`/dashboard/`）以關鍵字搜尋公告資料庫，後端使用 Django ORM 的 `Q` 物件實作全文搜尋（title、content 欄位），回傳結果以分頁方式呈現，可依來源或日期排序。')
add_figure_placeholder(doc, '【圖 4.9　公告儀表板搜尋結果頁面】')

# 4.8
add_section_heading(doc, '4.8　YouTube 影片情感儀表板（app_youtube_uma）')
add_subsection_heading(doc, '功能說明')
add_body(doc, '使用者於 YouTube 儀表板（`/youtube/`）查看系統自動抓取的賽馬娘相關 YouTube 影片列表，包含觀看數、按讚數、留言情感分數，以及週維度聲量趨勢折線圖。')
add_subsection_heading(doc, '執行過程')
for i, item in enumerate([
    'APScheduler 每 6 小時呼叫 YouTube Data API v3（每日 quota ≤ 3,000 units）',
    '搜尋關鍵字「賽馬娘」，擷取影片 metadata 及最新留言',
    '呼叫 Gemini API 對留言進行批次情感標記',
    '儲存至 `YouTubeVideo` 與 `YouTubeComment` 資料表',
], 1):
    add_numbered(doc, i, item)
add_figure_placeholder(doc, '【圖 4.10　YouTube 影片情感儀表板（影片卡片列表）】')
add_figure_placeholder(doc, '【圖 4.11　YouTube 週維度聲量與情感趨勢折線圖】')

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 第五章　Agentic AI 與 RAG 知識庫執行結果
# ══════════════════════════════════════════════════════
add_chapter_heading(doc, '第五章　Agentic AI 與 RAG 知識庫執行結果')

add_section_heading(doc, '5.1　Agentic AI 馬娘分析助理（app_agent_uma）')
add_subsection_heading(doc, '功能說明')
add_body(doc, '使用者於 Agentic AI 對話頁（`/agent/`）以自然語言提問，系統呼叫 Gemini 3.5 Flash，由 AI 自動選擇並呼叫適當工具，進行多輪推理後回傳完整分析結論。')
add_subsection_heading(doc, '工具函數（Function Calling）一覽')
add_table_caption(doc, '表 5.1　Agentic AI 工具函數一覽')
styled_table(doc,
    ['工具名稱', '功能說明'],
    [
        ['search_uma_news', '依關鍵字搜尋公告資料庫（NewsData ORM）'],
        ['get_character_popularity', '查詢指定角色的討論聲量'],
        ['get_sentiment_analysis', '取得關鍵字的情感分析結果'],
        ['get_top_keywords', '取得熱門關鍵詞排行（依類別）'],
        ['analyze_bahamut_comments', '分析巴哈姆特留言情緒分布'],
        ['read_local_document', '讀取 knowledge_base/ 目錄的本地知識文件'],
    ],
    col_widths=[6.0, 10.0]
)
add_subsection_heading(doc, '執行過程')
add_body(doc, '以問題「這週哪隻馬娘在巴哈討論最熱？」為例：')
for i, item in enumerate([
    '使用者輸入問題，前端 POST 至 `/agent/api/chat/`',
    '後端呼叫 Gemini API，附帶 6 個工具函數定義',
    'Gemini 判斷需呼叫 `get_character_popularity` 工具，後端執行 ORM 查詢',
    '查詢結果回傳 Gemini，Gemini 整合數據組合最終回答',
    '前端以 Markdown 格式顯示 AI 完整回應',
], 1):
    add_numbered(doc, i, item)
add_figure_placeholder(doc, '【圖 5.1　Agentic AI 對話頁面（多輪工具呼叫過程）】')
add_subsection_heading(doc, '執行成效')
for item in [
    '工具呼叫次數：每次問答平均 1–3 次工具呼叫',
    '回應時間：約 10–25 秒（視工具複雜度）',
    '支援多輪對話：可延續上文脈絡持續追問',
]:
    add_bullet(doc, item)

add_section_heading(doc, '5.2　LangChain ReAct Agent（app_agent_langchain）')
add_body(doc, '使用者於 LangChain Agent 頁面（`/langchain-agent/`）提問，系統採用 LangChain ReAct 框架，展示 Thought → Action → Observation 的完整推理迴圈。前端即時串流（Streaming）顯示 Agent 每一步的推理過程（思考鏈）。')
add_figure_placeholder(doc, '【圖 5.2　LangChain ReAct Agent 推理過程畫面（思考鏈可見）】')

add_section_heading(doc, '5.3　LangGraph 狀態圖 Agent（app_agent_langgraph）')
add_body(doc, '使用者於 LangGraph Agent 頁面（`/langgraph-agent/`）提問，系統採用 LangGraph StateGraph 管理 Agent 狀態轉換，展示複雜多步驟 Agentic AI 的最新框架。後端建立 StateGraph，定義節點（工具呼叫 / 回答生成）與邊（條件跳轉），Agent 依 StateGraph 執行並輸出完整回應。')
add_figure_placeholder(doc, '【圖 5.3　LangGraph StateGraph Agent 頁面（多步驟狀態轉換可視化）】')

add_section_heading(doc, '5.4　RAG 馬娘知識庫問答（app_rag_uma）')
add_subsection_heading(doc, '功能說明')
add_body(doc, '使用者於 RAG 問答頁面（`/rag/`）直接提問（預建索引常駐，無需上傳文件），或選擇上傳馬娘角色介紹 PDF / Markdown 追加知識庫，系統從 FAISS 向量索引中檢索相關段落並回傳含引用來源的精確答案。')
add_subsection_heading(doc, '執行過程')
add_body(doc, '**（一）索引建立（預建）**：`entrypoint.sh` 啟動時，讀取 `knowledge_base/` 目錄下的所有 `.md` 與 `.txt` 文件，呼叫 `gemini-embedding-001` 生成 768 維嵌入向量，儲存至 `app_rag_uma/index/uma_knowledge.faiss`。')
add_body(doc, '**（二）查詢流程**：使用者輸入問題後，後端對問題生成嵌入向量，FAISS 以 L2 距離搜尋最相近的 Top-5 文件段落，將段落與問題組合 Prompt，呼叫 Gemini 生成含引用標記的回答。')
add_body(doc, '**（三）持久化**：伺服器重啟後 FAISS 索引持續存在，無需重新建立。')
add_figure_placeholder(doc, '【圖 5.4　RAG 知識庫問答頁面（提問與引用來源顯示）】')
add_subsection_heading(doc, '執行成效')
for item in [
    '問題回答延遲：約 5–15 秒（含嵌入向量計算）',
    '知識庫持久化：✅ 重啟後可直接使用',
    '文件上傳功能：支援即時追加至現有索引',
]:
    add_bullet(doc, item)

add_section_heading(doc, '5.5　Agentic RAG 雙工具問答（app_rag_agent）')
add_body(doc, '使用者於 Agentic RAG 頁面（`/rag-agent/`）提問，Agent 根據問題性質自動選擇「語意搜尋 RAG 工具」或「資料庫精確查詢工具」，回答附引用新聞 ID 與連結，比傳統單一 RAG 更靈活。')
for item in [
    '若問題涉及「角色介紹、遊戲機制、技能效果」→ 優先使用 RAG 語意搜尋',
    '若問題涉及「特定公告、最新新聞、精確資料」→ 優先使用 ORM 資料庫查詢',
]:
    add_bullet(doc, item)
add_figure_placeholder(doc, '【圖 5.5　Agentic RAG 頁面（雙工具選擇邏輯與回答含引用 ID）】')

add_section_heading(doc, '5.6　全站 3D 虛擬客服「成田路」')
add_body(doc, '全站每一頁右下角常駐 3D 虛擬客服「馬娘客服 成田路」，透過 Three.js + VRM 格式的 3D 模型實現滑鼠跟隨視線、表情切換，並於使用者提問後以對話泡泡顯示 AI 回應（呼叫相同的 Agentic AI 後端），無需切換頁面即可互動。')
for i, item in enumerate([
    '使用者移動滑鼠時，角色頭部跟隨視線轉動',
    '使用者在輸入框輸入問題並送出',
    '角色顯示「思考中...」動畫',
    '後端呼叫 `/agent/api/chat/` 取得 AI 回覆',
    '角色顯示微笑表情，回覆以頭頂泡泡方式呈現',
], 1):
    add_numbered(doc, i, item)
add_figure_placeholder(doc, '【圖 5.6　3D 虛擬客服「成田路」即時互動畫面（滑鼠跟隨與泡泡回覆）】')

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 第六章　排程系統與 Discord Bot 執行結果
# ══════════════════════════════════════════════════════
add_chapter_heading(doc, '第六章　排程系統與 Discord Bot 執行結果')

add_section_heading(doc, '6.1　APScheduler 自動排程系統')
add_body(doc, '本平台整合 APScheduler 3.11.2 + django-apscheduler 0.7.0，於 Django 啟動時自動初始化排程，無需額外執行任何指令。')
add_table_caption(doc, '表 6.1　APScheduler 排程任務一覽')
styled_table(doc,
    ['任務名稱', '排程設定', '執行內容'],
    [
        ['巴哈留言爬取', '每 1 小時', 'scrape_bahamut management command'],
        ['留言情感分析', '每日 02:00', 'analyze_sentiment management command（呼叫 Gemini）'],
        ['YouTube 影片更新', '每 6 小時', '呼叫 YouTube Data API v3，更新影片與留言'],
        ['Discord 新聞推播', '每日 08:00', '爬取前日 Discord 頻道訊息，Gemini 彙整後推播'],
    ],
    col_widths=[4.0, 3.5, 8.5]
)
add_body(doc, '管理員於情報站控制台可查看排程狀態並手動觸發，點擊「立即執行」可觸發單次執行，或切換「啟動/停止排程」控制定期執行。')
add_figure_placeholder(doc, '【圖 6.1　APScheduler 排程管理頁面（任務列表與手動觸發按鈕）】')

add_section_heading(doc, '6.2　Discord Bot 自動新聞推播系統（D1-D8）')
add_body(doc, 'Discord Bot 為本平台的獨立服務容器，提供完整八項功能（D1-D8），如表 6.2 所示。')
add_table_caption(doc, '表 6.2　Discord Bot 功能列表')
styled_table(doc,
    ['功能代碼', '功能說明'],
    [
        ['D1', '啟動 / 停止 Discord Bot 服務'],
        ['D2', '設定監聽 Guild ID 與頻道 ID'],
        ['D3', '每 30 分鐘爬取頻道訊息'],
        ['D4', '第一層篩選：關鍵字過濾（馬娘相關）'],
        ['D5', '第二層篩選：Gemini 語意確認'],
        ['D6', '每日 08:00 由 Gemini 彙整成新聞稿'],
        ['D7', '自動推播新聞稿至指定 Discord 頻道'],
        ['D8', '管理頁（/discord/）查看推播歷史'],
    ],
    col_widths=[3.0, 13.0]
)
for i, item in enumerate([
    '管理員於情報站控制台的「Discord Bot 管理」頁設定 Bot Token 與目標頻道',
    '點擊「啟動 Bot」，Docker 容器中的 discord-bot 服務開始運行',
    'Bot 每 30 分鐘爬取監聽頻道內的新訊息',
    '雙層篩選過濾出馬娘相關訊息（關鍵字初篩 → Gemini 語意確認）',
    '每日 08:00 自動執行：Gemini 將前日篩選訊息彙整為繁體中文新聞稿',
    '新聞稿推播至指定 Discord 頻道，並記錄至 `DiscordNewsLog` 資料表',
], 1):
    add_numbered(doc, i, item)
add_figure_placeholder(doc, '【圖 6.2　Discord Bot 管理頁面（Bot 狀態、推播歷史、設定表單）】')
add_figure_placeholder(doc, '【圖 6.3　Discord 頻道中 Bot 自動推播的每日馬娘新聞稿】')

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 第七章　情報站控制台執行結果
# ══════════════════════════════════════════════════════
add_chapter_heading(doc, '第七章　情報站控制台執行結果')

add_section_heading(doc, '7.1　控制台總覽')
add_body(doc, '情報站控制台（`/crawler-admin/`）為本平台之管理中心，提供全平台健康監控與一站式操作入口。主頁儀表板顯示以下資訊：')
for item in [
    '**資料統計卡片**：各來源（Bahamut / Bilibili / ETtoday / UDN / Gamme）匯入筆數',
    '**YouTube API 配額進度條**：今日已使用 / 每日上限（10,000 units）',
    '**Discord Bot 狀態**：Online / Offline 指示燈、今日推播次數',
    '**RAG 索引狀態**：向量數量、索引建立時間',
    '**快速操作**：觸發爬蟲、重建 RAG 索引、手動 Discord 推播',
]:
    add_bullet(doc, item)
add_figure_placeholder(doc, '【圖 7.1　情報站控制台主頁儀表板（全平台健康狀態一覽）】')

add_section_heading(doc, '7.2　子功能頁面')
add_subsection_heading(doc, '（一）資料管理頁（Data Manager）')
add_body(doc, '顯示各來源的詳細統計，支援依條件清除特定來源資料，以及從 CSV 重新匯入。')
add_figure_placeholder(doc, '【圖 7.2　資料管理頁（各來源資料筆數與操作按鈕）】')
add_subsection_heading(doc, '（二）Pipeline 執行頁')
add_body(doc, '以分步驟可視化方式執行完整資料管線（爬取 → 前處理 → 情感標記 → 匯入 DB），每步驟完成後即時更新狀態。')
add_figure_placeholder(doc, '【圖 7.3　Pipeline 分步執行頁面（進度狀態顯示）】')
add_subsection_heading(doc, '（三）RAG 知識庫管理頁')
add_body(doc, '顯示 FAISS 索引向量數量、支援一鍵重建索引，以及上傳新知識文件。')
add_figure_placeholder(doc, '【圖 7.4　RAG 知識庫管理頁（索引狀態與重建操作）】')

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 第八章　系統部署與容器化執行結果
# ══════════════════════════════════════════════════════
add_chapter_heading(doc, '第八章　系統部署與容器化執行結果')

add_section_heading(doc, '8.1　Docker Compose 架構')
add_body(doc, '本平台採用 Docker Compose v2 進行容器化部署，包含三個服務容器，如表 8.1 所示。')
add_table_caption(doc, '表 8.1　Docker Compose 服務容器一覽')
styled_table(doc,
    ['服務名稱', '基礎映像', '說明'],
    [
        ['web-poa', 'python:3.12-slim-bookworm', 'Django + Gunicorn 主服務（Port 8000）'],
        ['discord-bot', 'python:3.12-slim-bookworm', 'Discord Bot 獨立服務'],
        ['nginx', 'nginx:1.26-alpine', '反向代理 + 靜態檔案服務（Port 80）'],
    ],
    col_widths=[3.5, 5.5, 7.0]
)
for item in [
    '`web-poa` 掛載 `entrypoint.sh` 為容器入口點，自動完成 migrate、collectstatic 與 Gunicorn 啟動',
    '`discord-bot` 與 `web-poa` 共用應用程式碼，透過環境變數指定啟動指令',
    '`nginx` 反向代理 `web-poa:8000`，並直接服務 `/static/` 目錄',
]:
    add_bullet(doc, item)

add_section_heading(doc, '8.2　部署執行結果')
add_body(doc, '執行 `docker compose up -d` 後，三個容器正常啟動，`docker compose ps` 輸出如下：')
add_code_block(doc, 'NAME              STATUS          PORTS\nnginx             Up              0.0.0.0:80->80/tcp\nweb-poa           Up              8000/tcp\ndiscord-bot       Up')
add_figure_placeholder(doc, '【圖 8.1　`docker compose ps` 輸出（三容器 Up 狀態）】')
add_figure_placeholder(doc, '【圖 8.2　瀏覽器透過 http://localhost 存取平台首頁（Nginx 服務確認）】')

add_section_heading(doc, '8.3　靜態檔案與媒體服務確認')
add_body(doc, 'Nginx 設定確認 `/static/` 路由正確對應至 `staticfiles/` 目錄：')
for item in [
    '前端 CSS / JS / 圖片資源：✅ 正常載入',
    'Bootstrap / Chart.js / Marked.js（CDN）：✅ 正常引入',
    '3D VRM 模型檔案（成田路）：✅ 正常載入',
]:
    add_bullet(doc, item)
add_figure_placeholder(doc, '【圖 8.3　瀏覽器開發者工具網路面板（靜態資源載入成功）】')

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 第九章　執行成效彙整
# ══════════════════════════════════════════════════════
add_chapter_heading(doc, '第九章　執行成效彙整')

add_section_heading(doc, '9.1　功能完成度彙整')
add_table_caption(doc, '表 9.1　課程需求完成度一覽')
styled_table(doc,
    ['課程要求', '需求內容', '對應功能', '完成狀態'],
    [
        ['要求 1', '以期中成果為基礎延伸', '以 w12 umamusume-llm-report 為基底', '✅ 完成'],
        ['要求 2', '自訂資料集完成基本 POA Apps', '5 來源資料集 + 全套基礎分析頁', '✅ 完成'],
        ['要求 3', 'Agentic AI', 'app_agent_uma（6 種 Function Calling）', '✅ 完成'],
        ['要求 4', 'RAG 或 Agentic RAG', 'app_rag_uma FAISS 持久化 + app_rag_agent 雙工具', '✅ 完成'],
        ['要求 5', '1–2 個 Feature App', '角色人氣 PK + 留言情感排程儀表板', '✅ 完成'],
        ['要求 6', '撰寫介紹頁面', 'app_poa_introduction（完整平台說明）', '✅ 完成'],
        ['要求 8', '大型語言模型資料分析', '雙模型（Gemini + Claude）分析報告', '✅ 完成'],
        ['部署', 'Docker Compose 容器化', 'Nginx + Gunicorn + Django 三層容器', '✅ 完成'],
        ['加分項', 'YouTube 社群資料', 'app_youtube_uma（YouTube Data API v3）', '✅ 完成'],
        ['加分項', 'LangChain/LangGraph Agent', 'app_agent_langchain + app_agent_langgraph', '✅ 完成'],
        ['加分項', 'Discord Bot 自動推播', 'D1-D8 完整實作', '✅ 完成'],
    ],
    col_widths=[2.5, 4.5, 6.5, 2.5]
)

add_section_heading(doc, '9.2　系統性能指標')
add_table_caption(doc, '表 9.2　核心功能執行時間彙整')
styled_table(doc,
    ['功能', '平均回應時間', '說明'],
    [
        ['首頁角色 PK 比較', '< 2 秒', 'ORM 查詢 + Chart.js 渲染'],
        ['關鍵詞聲量查詢', '< 1 秒', 'CSV 讀取 + Pandas 計算'],
        ['AI 雙模型分析報告（Gemini）', '15–25 秒', 'Gemini API 呼叫 + Markdown 渲染'],
        ['AI 雙模型分析報告（Claude）', '20–35 秒', 'Anthropic API 呼叫 + Markdown 渲染'],
        ['Agentic AI 問答', '10–25 秒', '1–3 次工具呼叫 + Gemini 推理'],
        ['RAG 知識庫問答', '5–15 秒', '嵌入向量計算 + FAISS 搜尋 + Gemini'],
        ['YouTube 儀表板載入', '< 1 秒', 'ORM 查詢（資料已預先分析）'],
    ],
    col_widths=[5.5, 3.5, 7.0]
)

add_section_heading(doc, '9.3　資料規模')
add_table_caption(doc, '表 9.3　資料庫資料規模（部署後）')
styled_table(doc,
    ['資料表', '資料模型', '資料來源', '說明'],
    [
        ['NewsData', '遊戲公告', '5 來源爬蟲', '含標題、內容、情感標籤'],
        ['Article', '巴哈貼文', '巴哈 API', '含貼文內容與情感六維度'],
        ['Comment', '巴哈留言', '巴哈 API', '對應貼文的留言資料'],
        ['YouTubeVideo', 'YouTube 影片', 'YouTube Data API v3', '含觀看數、按讚數'],
        ['YouTubeComment', 'YouTube 留言', 'YouTube Data API v3', '含情感分析結果'],
        ['DiscordMessage', 'Discord 訊息', 'Discord Bot 爬取', '篩選後的馬娘相關訊息'],
        ['DiscordNewsLog', '推播紀錄', 'Bot 自動生成', '每日推播新聞稿歷史'],
    ],
    col_widths=[4.0, 3.0, 4.5, 4.5]
)

add_section_heading(doc, '9.4　平台架構亮點')
for i, item in enumerate([
    '**多 Agent 框架並行實作**：本平台同時實作原生 Gemini Function Calling（`app_agent_uma`）、LangChain ReAct（`app_agent_langchain`）、LangGraph StateGraph（`app_agent_langgraph`）三種 Agentic AI 框架，提供橫向技術對照。',
    '**雙 LLM 對照分析**：AI 報告模組同時支援 Gemini 3.5 Flash 與 Claude Sonnet 4.6，使用者可直接對照兩個頂尖 LLM 的分析差異。',
    '**全站沉浸式互動**：3D VRM 成田路客服常駐全站，結合 Agentic AI 後端，提供傳統聊天室之上的角色陪伴體驗，具有高度視覺識別性。',
    '**外部服務深度整合**：YouTube 情感儀表板 + Discord Bot 自動推播，形成「社群監控 → AI 彙整 → 自動推播」的完整自動化輿情循環。',
    '**容器化一鍵部署**：`docker compose up -d` 搭配 idempotent 的 `entrypoint.sh`，確保每次部署均可重現一致的執行狀態。',
], 1):
    add_numbered(doc, i, item)

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 附錄
# ══════════════════════════════════════════════════════
add_chapter_heading(doc, '附錄 A　目錄結構一覽')
add_code_block(doc,
"""umamusume-information-platform/
├── app_agent_langchain/        # LangChain ReAct Agent
├── app_agent_langgraph/        # LangGraph StateGraph Agent
├── app_agent_uma/              # Agentic AI (Function Calling)
├── app_character_pk/           # 首頁角色人氣 PK
├── app_comment_sentiment/      # 留言情感排程儀表板
├── app_correlation_analysis/   # 關鍵詞相關性分析
├── app_course_intro/           # 課程技術說明頁
├── app_crawler_admin/          # 情報站控制台
├── app_dashboard/              # 公告列表儀表板
├── app_discord_bot/            # Discord Bot 自動推播
├── app_poa_introduction/       # 平台介紹頁
├── app_rag_agent/              # Agentic RAG（雙工具）
├── app_rag_uma/                # RAG + FAISS 知識庫
├── app_uma_comments/           # 巴哈留言資料模型
├── app_uma_news/               # 遊戲公告資料模型
├── app_uma_top_character/      # 熱門角色排行
├── app_uma_top_keyword/        # 熱門關鍵詞排行
├── app_user_keyword/           # 關鍵詞聲量分析（CSV）
├── app_user_keyword_llm_report/ # 雙模型 AI 報告
├── app_user_keyword_sentiment/ # 關鍵詞情感分析
├── app_youtube_uma/            # YouTube 影片情感儀表板
├── data/raw/                   # 各爬蟲原始 CSV
├── data/processed/             # 前處理與情感標記後 CSV
├── knowledge_base/             # RAG + Agent 知識文件
├── pipeline/                   # 資料管線腳本
├── scripts/                    # 資料匯入與統計腳本
├── website_configs/            # Django settings + urls
├── docker-compose.yml          # 容器編排設定
├── nginx/nginx.conf            # Nginx 反向代理設定
└── requirements.txt            # Python 相依套件""")

add_chapter_heading(doc, '附錄 B　主要 API 端點一覽')
add_table_caption(doc, '表 B.1　核心 API 端點')
styled_table(doc,
    ['端點', 'HTTP 方法', '功能'],
    [
        ['/agent/api/chat/', 'POST', 'Agentic AI 對話（含 Function Calling）'],
        ['/rag/api/query/', 'POST', 'RAG 知識庫問答'],
        ['/langchain-agent/api/chat/', 'POST', 'LangChain ReAct Agent 對話'],
        ['/langgraph-agent/api/chat/', 'POST', 'LangGraph StateGraph Agent 對話'],
        ['/llm-report/api/generate/', 'POST', '雙模型 LLM 分析報告生成'],
        ['/youtube/api/stats/', 'GET', 'YouTube 週趨勢情感統計'],
        ['/crawler-admin/api/run-pipeline/', 'POST', '手動觸發資料管線'],
        ['/crawler-admin/api/rebuild-rag/', 'POST', '重建 RAG FAISS 索引'],
        ['/crawler-admin/api/discord/messages/', 'GET', '取得 Discord 推播歷史'],
    ],
    col_widths=[6.5, 3.0, 6.5]
)

add_chapter_heading(doc, '附錄 C　環境變數說明')
add_table_caption(doc, '表 C.1　.env 環境變數一覽')
styled_table(doc,
    ['變數名稱', '說明', '必填'],
    [
        ['GEMINI_API_KEY', 'Google Gemini API 金鑰', '✅'],
        ['ANTHROPIC_API_KEY', 'Anthropic Claude API 金鑰', '✅'],
        ['YOUTUBE_API_KEY', 'YouTube Data API v3 金鑰', '✅（O5 功能）'],
        ['DISCORD_BOT_TOKEN', 'Discord Bot Token', '✅（D1-D8 功能）'],
        ['UMA_CHAT_MODEL', 'Gemini 對話模型名稱（預設 gemini-3.5-flash）', '❌'],
        ['SECRET_KEY', 'Django Secret Key', '✅'],
        ['DEBUG', '除錯模式（True/False）', '❌（預設 True）'],
    ],
    col_widths=[5.0, 9.0, 2.0]
)

# 尾部資訊
p = doc.add_paragraph()
set_para_spacing(p, 20, 0)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('─' * 30)
r.font.color.rgb = C_GOLD

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p, 6, 0)
for line in [
    '本報告撰寫日期：2026-06-25',
    '平台版本：v0.6.51-alpha',
    '技術棧：Django 5.2 LTS + Gemini 3.5 Flash + Claude Sonnet 4.6 + FAISS + Docker Compose',
]:
    r = p.add_run(line + '\n')
    r.font.name = '標楷體'
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ══════════════════════════════════════════════════════
# 儲存
# ══════════════════════════════════════════════════════
output_path = r'c:\Users\admin\dev\umamusume-information-platform\論文_程式執行過程與結果.docx'
doc.save(output_path)
print(f'OK saved: {output_path}')
